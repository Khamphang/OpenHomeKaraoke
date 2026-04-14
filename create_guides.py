"""Generate User Guide and Testing Guide for OpenHomeKaraoke in Lao language."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear'
    })
    shading.append(shading_elem)

def add_styled_table(doc, headers, rows, header_color="1a5276"):
    """Add a styled table with colored header."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(11)
        set_cell_shading(cell, header_color)

    # Data rows
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
            if r % 2 == 1:
                set_cell_shading(cell, "ebf5fb")

    return table

def add_title_page(doc, title, subtitle):
    """Add a title page."""
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1a, 0x52, 0x76)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(subtitle)
    run2.font.size = Pt(16)
    run2.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("OpenHomeKaraoke v2.5.0")
    run3.font.size = Pt(14)
    run3.font.color.rgb = RGBColor(0x85, 0x92, 0x9E)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run("ວັນທີ: ເມສາ 2026")
    run4.font.size = Pt(12)
    run4.font.color.rgb = RGBColor(0x85, 0x92, 0x9E)

    doc.add_page_break()


def create_user_guide():
    """Create the User Guide document."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    add_title_page(doc,
        "OpenHomeKaraoke",
        "ຄູ່ມືການນຳໃຊ້ (User Guide)"
    )

    # ===== ສາລະບານ =====
    doc.add_heading("ສາລະບານ", level=1)
    toc_items = [
        "1. ແນະນຳທົ່ວໄປ",
        "2. ຄວາມຕ້ອງການຂອງລະບົບ",
        "3. ການຕິດຕັ້ງ",
        "4. ການເປີດໃຊ້ງານ",
        "5. ໜ້າຫຼັກ (Home) - ຄວບຄຸມການຫຼິ້ນເພງ",
        "6. ໜ້າຄິວ (Queue) - ຈັດການຄິວເພງ",
        "7. ໜ້າຄົ້ນຫາ (Search) - ຊອກຫາ ແລະ ດາວໂຫຼດເພງ",
        "8. ໜ້າເບິ່ງເພງ (Browse) - ລາຍການເພງທັງໝົດ",
        "9. ໜ້າຂໍ້ມູນ (Info) - ຕັ້ງຄ່າ ແລະ ຂໍ້ມູນລະບົບ",
        "10. ການໃຊ້ງານຜ່ານມືຖື",
        "11. ຟີເຈີ Admin",
        "12. ຄຳຖາມທີ່ພົບເລື້ອຍ (FAQ)",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ===== 1. ແນະນຳທົ່ວໄປ =====
    doc.add_heading("1. ແນະນຳທົ່ວໄປ", level=1)
    doc.add_paragraph(
        "OpenHomeKaraoke ແມ່ນລະບົບຄາລາໂອເກະແບບ open-source ທີ່ສາມາດ"
        "ຄົ້ນຫາ ແລະ ດາວໂຫຼດເພງຈາກ YouTube ໄດ້ໂດຍກົງ. "
        "ຜູ້ໃຊ້ສາມາດເຂົ້າໃຊ້ຜ່ານ browser ທັງໃນຄອມພິວເຕີ ແລະ ມືຖືໃນ WiFi ດຽວກັນ."
    )

    doc.add_heading("ຄຸນສົມບັດຫຼັກ", level=2)
    features = [
        "ຄົ້ນຫາ ແລະ ດາວໂຫຼດເພງຈາກ YouTube",
        "ຈັດຄິວເພງ ແລະ ເປີດເພງແບບ real-time",
        "ຄວບຄຸມ volume, ຄວາມໄວ, ແລະ key ເພງ",
        "ແຍກສຽງຮ້ອງ/ດົນຕີ (Vocal Splitter)",
        "ຮອງຮັບຫຼາຍພາສາ",
        "ເຂົ້າໃຊ້ຜ່ານມືຖືດ້ວຍ QR Code",
        "ລະບົບ Admin ປ້ອງກັນດ້ວຍລະຫັດຜ່ານ",
    ]
    for f in features:
        doc.add_paragraph(f, style='List Bullet')

    # ===== 2. ຄວາມຕ້ອງການ =====
    doc.add_heading("2. ຄວາມຕ້ອງການຂອງລະບົບ", level=1)
    add_styled_table(doc,
        ["ລາຍການ", "ລາຍລະອຽດ"],
        [
            ["ລະບົບປະຕິບັດການ", "Windows 10/11, macOS, Linux, Raspberry Pi"],
            ["Python", "3.8 ຂຶ້ນໄປ (ແນະນຳ 3.11)"],
            ["VLC Media Player", "ເວີຊັນ 3.x (ຕິດຕັ້ງໃສ່ path ເລີ່ມຕົ້ນ)"],
            ["RAM", "4 GB ຂຶ້ນໄປ"],
            ["ອິນເຕີເນັດ", "ຕ້ອງການສຳລັບດາວໂຫຼດເພງຈາກ YouTube"],
            ["Browser", "Chrome, Firefox, Edge, Safari"],
        ]
    )

    # ===== 3. ການຕິດຕັ້ງ =====
    doc.add_heading("3. ການຕິດຕັ້ງ", level=1)

    doc.add_heading("ຂັ້ນຕອນ 1: ຕິດຕັ້ງ VLC", level=2)
    doc.add_paragraph("ເປີດ Terminal/Command Prompt ແລ້ວພິມ:")
    p = doc.add_paragraph()
    run = p.add_run("winget install VideoLAN.VLC")
    run.font.name = "Consolas"
    run.font.size = Pt(10)

    doc.add_heading("ຂັ້ນຕອນ 2: Clone ໂປຣເຈັກ", level=2)
    p = doc.add_paragraph()
    run = p.add_run("git clone https://github.com/xuancong84/OpenHomeKaraoke.git")
    run.font.name = "Consolas"
    run.font.size = Pt(10)

    doc.add_heading("ຂັ້ນຕອນ 3: ຕິດຕັ້ງ dependencies", level=2)
    p = doc.add_paragraph()
    run = p.add_run("cd OpenHomeKaraoke\npip install -r requirements.txt")
    run.font.name = "Consolas"
    run.font.size = Pt(10)

    doc.add_heading("ຂັ້ນຕອນ 4: ສ້າງ folder ເພງ", level=2)
    p = doc.add_paragraph()
    run = p.add_run("mkdir songs\\vocal songs\\nonvocal")
    run.font.name = "Consolas"
    run.font.size = Pt(10)

    # ===== 4. ການເປີດໃຊ້ງານ =====
    doc.add_heading("4. ການເປີດໃຊ້ງານ", level=1)

    doc.add_heading("ຄຳສັ່ງພື້ນຖານ", level=2)
    p = doc.add_paragraph()
    run = p.add_run("python app.py -d songs")
    run.font.name = "Consolas"
    run.font.size = Pt(10)

    doc.add_paragraph(
        "ຫຼັງຈາກແລ່ນແລ້ວ, ເປີດ browser ໄປທີ່ http://127.0.0.1:5000"
    )

    doc.add_heading("ຄຳສັ່ງເພີ່ມເຕີມ", level=2)
    add_styled_table(doc,
        ["ຄຳສັ່ງ", "ຄຳອະທິບາຍ"],
        [
            ["-p 8080", "ປ່ຽນ port ເປັນ 8080"],
            ["--ssl", "ເປີດ HTTPS (ຈຳເປັນສຳລັບໃຊ້ microphone)"],
            ["-w", "ເປີດແບບ windowed (ບໍ່ fullscreen)"],
            ["-hq", "ດາວໂຫຼດເພງຄຸນນະພາບສູງ"],
            ["--admin-password xxxx", "ຕັ້ງລະຫັດ admin"],
            ["-L lo", "ຕັ້ງພາສາເປັນລາວ"],
            ["-s 5", "ລໍຖ້າ 5 ວິ ລະຫວ່າງເພງ"],
            ["-nv", "ເປີດການປັບ volume ອັດຕະໂນມັດ"],
        ]
    )

    doc.add_paragraph()
    p = doc.add_paragraph("ຕົວຢ່າງ: ")
    run = p.add_run('python app.py -d songs -p 8080 --ssl --admin-password mypass123 -w')
    run.font.name = "Consolas"
    run.font.size = Pt(10)

    doc.add_page_break()

    # ===== 5. ໜ້າຫຼັກ (Home) =====
    doc.add_heading("5. ໜ້າຫຼັກ (Home) - ຄວບຄຸມການຫຼິ້ນເພງ", level=1)
    doc.add_paragraph(
        "ໜ້ານີ້ແມ່ນໜ້າຄວບຄຸມຫຼັກ ສະແດງເພງທີ່ກຳລັງຫຼິ້ນ ແລະ ເພງຕໍ່ໄປໃນຄິວ."
    )

    doc.add_heading("ປຸ່ມຄວບຄຸມການຫຼິ້ນ", level=2)
    add_styled_table(doc,
        ["ປຸ່ມ", "ໜ້າທີ່"],
        [
            ["Restart", "ເລີ່ມເພງໃໝ່ຈາກຕົ້ນ"],
            ["Play/Pause", "ຢຸດຊົ່ວຄາວ / ຫຼິ້ນຕໍ່"],
            ["Skip", "ຂ້າಮເພງ, ໄປເພງຕໍ່ໄປ"],
            ["Volume -/+", "ປັບສຽງດັງ-ເບົາ"],
            ["Seek Bar", "ເລື່ອນໄປຕຳແໜ່ງທີ່ຕ້ອງການ"],
        ]
    )

    doc.add_heading("ການຄວບຄຸມສຽງຂັ້ນສູງ", level=2)
    advanced_audio = [
        "Audio Delay: ປັບຄວາມຊ້າ-ໄວ ຂອງສຽງ (ວິນາທີ)",
        "Subtitle Delay: ປັບຄວາມຊ້າ-ໄວ ຂອງ subtitle",
        "Playback Speed: ປັບຄວາມໄວ 0.8x - 1.2x",
        "Key/Transpose: ປ່ຽນ key ເພງ -12 ຫາ +12 semitones (VLC ເທົ່ານັ້ນ)",
    ]
    for item in advanced_audio:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("ໂໝດສຽງຮ້ອງ (Vocal Mode)", level=2)
    add_styled_table(doc,
        ["ໂໝດ", "ຄຳອະທິບາຍ"],
        [
            ["Nonvocal", "ສະແດງສະເພາະດົນຕີ (ສຳລັບຮ້ອງເອງ)"],
            ["Mixed", "ສຽງຮ້ອງ + ດົນຕີ (ຕົ້ນສະບັບ)"],
            ["Vocal", "ສະເພາະສຽງຮ້ອງ (ສຳລັບຝຶກ)"],
        ]
    )

    doc.add_page_break()

    # ===== 6. ໜ້າຄິວ (Queue) =====
    doc.add_heading("6. ໜ້າຄິວ (Queue) - ຈັດການຄິວເພງ", level=1)
    doc.add_paragraph(
        "ໜ້ານີ້ສະແດງລາຍການເພງທີ່ລໍຖ້າ ແລະ ອະນຸຍາດໃຫ້ຈັດລຳດັບຄິວ."
    )

    doc.add_heading("ການຈັດການຄິວ", level=2)
    queue_features = [
        "ເບິ່ງລາຍການເພງທັງໝົດໃນຄິວ ພ້ອມຊື່ຜູ້ຮ້ອງ",
        "ລາກ-ວາງ (Drag & Drop) ເພື່ອປ່ຽນລຳດັບ (ເທິງ desktop)",
        "ປຸ່ມຍ້າຍຂຶ້ນ/ລົງ ເພື່ອປ່ຽນລຳດັບ",
        "ປຸ່ມລຶບ ເພື່ອເອົາເພງອອກຈາກຄິວ",
        "ປຸ່ມ 'ເພີ່ມ 3 ເພງແບບສຸ່ມ' ເພື່ອເຕີມເພງອັດຕະໂນມັດ",
        "ປຸ່ມ 'ລ້າງທັງໝົດ' ເພື່ອລຶບທຸກເພງໃນຄິວ (ຕ້ອງຢືນຢັນ)",
    ]
    for item in queue_features:
        doc.add_paragraph(item, style='List Bullet')

    # ===== 7. ໜ້າຄົ້ນຫາ (Search) =====
    doc.add_heading("7. ໜ້າຄົ້ນຫາ (Search) - ຊອກຫາ ແລະ ດາວໂຫຼດເພງ", level=1)

    doc.add_heading("ການຄົ້ນຫາເພງ", level=2)
    doc.add_paragraph(
        "ພິມຊື່ເພງ ຫຼື ສິລະປິນ ໃນຊ່ອງຄົ້ນຫາ. "
        "ລະບົບຈະສະແດງຜົນຈາກເພງທີ່ມີຢູ່ແລ້ວ ແລະ ຈາກ YouTube."
    )

    search_steps = [
        "ພິມຊື່ເພງໃນຊ່ອງ Search",
        "ຖ້າເພງມີຢູ່ແລ້ວ → ກົດ 'Add to Queue' ເພື່ອເພີ່ມເຂົ້າຄິວ",
        "ຖ້າຕ້ອງການຈາກ YouTube → ກົດ Search ແລ້ວເລືອກຈາກ 10 ຜົນລັບ",
        "ເລືອກ options: Enqueue (ເພີ່ມຄິວ), High Quality, Include Subtitles",
        "ກົດ Download ເພື່ອດາວໂຫຼດ",
    ]
    for i, step in enumerate(search_steps, 1):
        doc.add_paragraph(f"{i}. {step}")

    doc.add_heading("ຄົ້ນຫາດ້ວຍສຽງ (Voice Search)", level=2)
    doc.add_paragraph(
        "ກົດປຸ່ມ microphone ເພື່ອເວົ້າຊື່ເພງ (ຕ້ອງເປີດ HTTPS ກ່ອນ). "
        "ລະບົບຈະແປງສຽງເປັນຂໍ້ຄວາມ ແລະ ຄົ້ນຫາໃຫ້ອັດຕະໂນມັດ."
    )

    doc.add_heading("ດາວໂຫຼດຈາກ URL ໂດຍກົງ", level=2)
    doc.add_paragraph(
        "ເປີດ Advanced Settings → ວາງ URL ຂອງ YouTube → ກົດ Download."
    )

    doc.add_page_break()

    # ===== 8. ໜ້າເບິ່ງເພງ (Browse) =====
    doc.add_heading("8. ໜ້າເບິ່ງເພງ (Browse) - ລາຍການເພງທັງໝົດ", level=1)
    doc.add_paragraph(
        "ສະແດງເພງທັງໝົດທີ່ດາວໂຫຼດໄວ້ແລ້ວ. ສາມາດຈັດລຽງ ແລະ ກັ່ນຕອງໄດ້."
    )

    browse_features = [
        "ຈັດລຽງຕາມ: ຕົວອັກສອນ (A-Z) ຫຼື ວັນທີ",
        "ແຖບຕົວອັກສອນ (A-Z, #) ເພື່ອກົດໄປຫາເພງໄວ",
        "ກົດ + ເພື່ອເພີ່ມເພງເຂົ້າຄິວ",
        "Admin ສາມາດແກ້ໄຂຊື່ ຫຼື ລຶບເພງໄດ້",
        "ແບ່ງໜ້າ 500 ເພງຕໍ່ໜ້າ",
    ]
    for item in browse_features:
        doc.add_paragraph(item, style='List Bullet')

    # ===== 9. ໜ້າຂໍ້ມູນ (Info) =====
    doc.add_heading("9. ໜ້າຂໍ້ມູນ (Info) - ຕັ້ງຄ່າ ແລະ ຂໍ້ມູນລະບົບ", level=1)

    doc.add_heading("ຂໍ້ມູນລະບົບ", level=2)
    info_items = [
        "URL ແລະ QR Code ສຳລັບເຊື່ອಮຕໍ່",
        "ຈຳນວນເພງທັງໝົດ",
        "CPU, RAM, Disk ທີ່ໃຊ້",
        "ເວີຊັນ yt-dlp ແລະ app",
    ]
    for item in info_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("ການຕັ້ງຄ່າ", level=2)
    add_styled_table(doc,
        ["ການຕັ້ງຄ່າ", "ຄຳອະທິບາຍ"],
        [
            ["Save Delays", "ບັນທຶກ audio/subtitle delay ຕໍ່ເພງ"],
            ["Volume Normalization", "ປັບ volume ໃຫ້ເທົ່າກັນທຸກເພງ"],
            ["Vocal Splitter Mode", "ເລືອກ DNN ຫຼື Traditional"],
            ["ປ່ຽນພາສາ", "ເລືອກພາສາ interface"],
            ["Refresh Database", "ສະແກນເພງໃໝ່"],
            ["Update yt-dlp", "ອັບເດດ YouTube downloader"],
        ]
    )

    doc.add_page_break()

    # ===== 10. ການໃຊ້ງານຜ່ານມືຖື =====
    doc.add_heading("10. ການໃຊ້ງານຜ່ານມືຖື", level=1)

    mobile_steps = [
        "ໃຫ້ມືຖືເຊື່ອม WiFi ດຽວກັນກັບເຄື່ອງທີ່ແລ່ນ app",
        "ສະແກນ QR Code ທີ່ສະແດງໃນໜ້າ Info ຫຼື splash screen",
        "ຫຼື ພິມ IP address ໃນ browser ເຊັ່ນ: http://192.168.1.22:5000",
        "ໃສ່ຊື່ຜູ້ຮ້ອງ (ຈະຖືກບັນທຶກໄວ້ໃນ cookie)",
        "ຄົ້ນຫາເພງ ແລະ ເພີ່ມເຂົ້າຄິວໄດ້ເລີຍ",
    ]
    for i, step in enumerate(mobile_steps, 1):
        doc.add_paragraph(f"{i}. {step}")

    # ===== 11. ຟີເຈີ Admin =====
    doc.add_heading("11. ຟີເຈີ Admin", level=1)
    doc.add_paragraph(
        "ຖ້າຕັ້ງ --admin-password, ຟີເຈີບາງອັນຕ້ອງ login ກ່ອນ:"
    )

    admin_features = [
        "ຍ້າຍ/ລຶບເພງໃນຄິວ",
        "ຂ້าม/ຢຸດ/ຫຼິ້ນເພງ",
        "ແກ້ໄຂ/ລຶບເພງຈາກ library",
        "ລ້າງຄິວທັງໝົດ",
        "Reboot/Shutdown ເຄື່ອງ (Raspberry Pi)",
        "ຕັ້ງຄ່າ vocal splitter ແລະ volume normalization",
    ]
    for item in admin_features:
        doc.add_paragraph(item, style='List Bullet')

    # ===== 12. FAQ =====
    doc.add_heading("12. ຄຳຖາມທີ່ພົບເລື້ອຍ (FAQ)", level=1)

    faqs = [
        ("ເພງບໍ່ມີສຽງ?", "ກວດ VLC ວ່າຕິດຕັ້ງຖືກບ່ອນ. ກວດ volume ໃນ app ແລະ ໃນ Windows."),
        ("ດາວໂຫຼດເພງບໍ່ໄດ້?", "ກວດອິນເຕີເນັດ. ລອງ Update yt-dlp ໃນໜ້າ Info."),
        ("ມືຖືເຂົ້າບໍ່ໄດ້?", "ກວດ WiFi ດຽວກັນ. ກວດ Firewall ບໍ່ block port 5000."),
        ("Voice search ບໍ່ເຮັດວຽກ?", "ຕ້ອງເປີດ --ssl ແລະ ຕ້ອງມີ --cloud server."),
        ("ຢາກປ່ຽນ key ເພງ?", "ໃຊ້ Key/Transpose slider ໃນໜ້າ Home (-12 ຫາ +12)."),
    ]
    for q, a in faqs:
        p = doc.add_paragraph()
        run_q = p.add_run(f"ຖ: {q}")
        run_q.bold = True
        p2 = doc.add_paragraph(f"ຕ: {a}")
        p2.paragraph_format.space_after = Pt(8)

    # Save
    output_path = os.path.join(os.path.dirname(__file__), "Guide_OpenHomeKaraoke.docx")
    doc.save(output_path)
    print(f"User Guide saved: {output_path}")
    return output_path


def create_test_guide():
    """Create the Testing Guide document."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    add_title_page(doc,
        "OpenHomeKaraoke",
        "ຄູ່ມືການທົດສອບ (Testing Guide)"
    )

    # ===== ສາລະບານ =====
    doc.add_heading("ສາລະບານ", level=1)
    toc_items = [
        "1. ຈຸດປະສົງ ແລະ ຂອບເຂດການທົດສອບ",
        "2. ສະພາບແວດລ້ອມ ແລະ ເງື່ອນໄຂກ່ອນທົດສອບ",
        "3. TC-01: ການຕິດຕັ້ງ ແລະ ເປີດ App",
        "4. TC-02: ໜ້າຫຼັກ (Home) - ຄວບຄຸມການຫຼິ້ນ",
        "5. TC-03: ຄົ້ນຫາ ແລະ ດາວໂຫຼດເພງ",
        "6. TC-04: ຈັດການຄິວເພງ",
        "7. TC-05: ເບິ່ງ ແລະ ຈັດການເພງ (Browse)",
        "8. TC-06: ການໃຊ້ງານຜ່ານມືຖື",
        "9. TC-07: ລະບົບ Admin",
        "10. TC-08: ໜ້າ Info ແລະ ການຕັ້ງຄ່າ",
        "11. TC-09: ການທົດສອບ Edge Cases",
        "12. ສະຫຼຸບ ແລະ ເກນການຜ່ານ",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ===== 1. ຈຸດປະສົງ =====
    doc.add_heading("1. ຈຸດປະສົງ ແລະ ຂອບເຂດການທົດສອບ", level=1)
    doc.add_paragraph(
        "ເອກະສານນີ້ກຳນົດ test cases ສຳລັບທົດສອບ OpenHomeKaraoke v2.5.0 "
        "ເພື່ອຮັບປະກັນວ່າທຸກຟີເຈີເຮັດວຽກຖືກຕ້ອງກ່ອນນຳໄປໃຊ້ໃນງານຈິງ."
    )

    doc.add_heading("ປະເພດການທົດສອບ", level=2)
    test_types = [
        "Functional Testing - ທົດສອບທຸກຟີເຈີເຮັດວຽກຖືກຕ້ອງ",
        "UI Testing - ກວດ interface ສະແດງຜົນຖືກຕ້ອງ",
        "Mobile Testing - ທົດສອບການໃຊ້ງານຜ່ານມືຖື",
        "Admin Testing - ທົດສອບ authentication ແລະ ສິດການເຂົ້າເຖິງ",
        "Edge Case Testing - ທົດສອບກໍລະນີພິເສດ",
    ]
    for t in test_types:
        doc.add_paragraph(t, style='List Bullet')

    # ===== 2. ເງື່ອນໄຂກ່ອນ =====
    doc.add_heading("2. ສະພາບແວດລ້ອມ ແລະ ເງື່ອນໄຂກ່ອນທົດສອບ", level=1)
    add_styled_table(doc,
        ["ລາຍການ", "ລາຍລະອຽດ"],
        [
            ["ເຄື່ອງທົດສອບ", "Windows 11, Python 3.11, VLC 3.0.23"],
            ["Browser", "Chrome (ລ່າສຸດ), Firefox (ລ່າສຸດ)"],
            ["ມືຖື", "Android / iOS ກັບ Chrome / Safari"],
            ["ອິນເຕີເນັດ", "ຕ້ອງມີ (ສຳລັບ YouTube)"],
            ["ຄຳສັ່ງແລ່ນ", "py -3.11 app.py -d songs -w"],
            ["URL ທົດສອບ", "http://127.0.0.1:5000"],
        ]
    )

    doc.add_page_break()

    # ===== TC-01 =====
    doc.add_heading("3. TC-01: ການຕິດຕັ້ງ ແລະ ເປີດ App", level=1)
    add_styled_table(doc,
        ["ID", "ລາຍການທົດສອບ", "ຂັ້ນຕອນ", "ຜົນທີ່ຄາດຫວັງ", "ຜົນຈິງ", "ຜ່ານ?"],
        [
            ["1.1", "ເປີດ app ດ້ວຍຄຳສັ່ງພື້ນຖານ",
             "ແລ່ນ: py -3.11 app.py -d songs",
             "App ເລີ່ມ, ສະແດງ URL http://127.0.0.1:5000", "", ""],
            ["1.2", "ເປີດ browser",
             "ພິມ http://127.0.0.1:5000",
             "ສະແດງໜ້າ Home ຂອງ OpenHomeKaraoke", "", ""],
            ["1.3", "ເປີດແບບ windowed",
             "ແລ່ນ: py -3.11 app.py -d songs -w",
             "VLC ເປີດແບບ window ບໍ່ fullscreen", "", ""],
            ["1.4", "ເປີດກັບ port ອື່ນ",
             "ແລ່ນ: py -3.11 app.py -d songs -p 8080",
             "App ແລ່ນທີ່ port 8080", "", ""],
            ["1.5", "ເປີດກັບ admin password",
             "ແລ່ນ: py -3.11 app.py -d songs --admin-password test123",
             "App ເລີ່ມ, ຟີເຈີ admin ຕ້ອງ login", "", ""],
        ]
    )

    doc.add_page_break()

    # ===== TC-02 =====
    doc.add_heading("4. TC-02: ໜ້າຫຼັກ (Home) - ຄວບຄຸມການຫຼິ້ນ", level=1)
    add_styled_table(doc,
        ["ID", "ລາຍການທົດສອບ", "ຂັ້ນຕອນ", "ຜົນທີ່ຄາດຫວັງ", "ຜົນຈິງ", "ຜ່ານ?"],
        [
            ["2.1", "ສະແດງເພງປັດຈຸບັນ",
             "ເພີ່ມເພງໃນຄິວ → ໄປໜ້າ Home",
             "ສະແດງຊື່ເພງ ແລະ ຜູ້ຮ້ອງ", "", ""],
            ["2.2", "Play/Pause",
             "ກົດ Play/Pause",
             "ເພງຢຸດ/ຫຼິ້ນຕໍ່", "", ""],
            ["2.3", "Skip ເພງ",
             "ກົດ Skip",
             "ຂ້າມໄປເພງຕໍ່ໄປ", "", ""],
            ["2.4", "Restart ເພງ",
             "ກົດ Restart",
             "ເພງເລີ່ມໃໝ່ຈາກ 0:00", "", ""],
            ["2.5", "ປັບ Volume",
             "ກົດ Volume + ແລະ -",
             "ສຽງດັງຂຶ້ນ/ເບົາລົງ", "", ""],
            ["2.6", "Seek bar",
             "ລາກ seek bar ໄປກາງເພງ",
             "ເພງໂດດໄປຕຳແໜ່ງທີ່ເລືອກ", "", ""],
            ["2.7", "ປັບ Playback Speed",
             "ປ່ຽນ speed ເປັນ 0.8x ແລະ 1.2x",
             "ເພງຊ້າລົງ/ໄວຂຶ້ນ", "", ""],
            ["2.8", "ປ່ຽນ Key/Transpose",
             "ລາກ slider ໄປ +3",
             "Key ເພງສູງຂຶ້ນ 3 semitones", "", ""],
            ["2.9", "ປ່ຽນ Vocal Mode",
             "ກົດ Nonvocal → Mixed → Vocal",
             "ສຽງປ່ຽນຕາມ mode", "", ""],
            ["2.10", "Audio Delay",
             "ປັບ Audio Delay +0.5",
             "ສຽງຊ້າລົງ 0.5 ວິ", "", ""],
        ]
    )

    doc.add_page_break()

    # ===== TC-03 =====
    doc.add_heading("5. TC-03: ຄົ້ນຫາ ແລະ ດາວໂຫຼດເພງ", level=1)
    add_styled_table(doc,
        ["ID", "ລາຍການທົດສອບ", "ຂັ້ນຕອນ", "ຜົນທີ່ຄາດຫວັງ", "ຜົນຈິງ", "ຜ່ານ?"],
        [
            ["3.1", "ຄົ້ນຫາເພງທີ່ມີຢູ່ແລ້ວ",
             "ພິມຊື່ເພງທີ່ download ແລ້ວ",
             "ສະແດງ autocomplete + ປຸ່ມ Add to Queue", "", ""],
            ["3.2", "ຄົ້ນຫາເພງຈາກ YouTube",
             "ພິມຊື່ເພງ → ກົດ Search",
             "ສະແດງ 10 ຜົນລັບຈາກ YouTube", "", ""],
            ["3.3", "ດາວໂຫຼດເພງ",
             "ເລືອກເພງ → ກົດ Download",
             "ເພງດາວໂຫຼດສຳເລັດ, ສະແດງໃນ Browse", "", ""],
            ["3.4", "ດາວໂຫຼດ + Enqueue",
             "ເລືອກ Enqueue → ກົດ Download",
             "ເພງດາວໂຫຼດ ແລະ ເພີ່ມເຂົ້າຄິວອັດຕະໂນມັດ", "", ""],
            ["3.5", "ດາວໂຫຼດ High Quality",
             "ເລືອກ High Quality → Download",
             "ເພງດາວໂຫຼດຄຸນນະພາບສູງກວ່າ", "", ""],
            ["3.6", "ດາວໂຫຼດຈາກ URL ໂດຍກົງ",
             "ວາງ YouTube URL → Download",
             "ເພງດາວໂຫຼດສຳເລັດ", "", ""],
            ["3.7", "ຄົ້ນຫາ Karaoke ເທົ່ານັ້ນ",
             "ເລືອກ 'Search Karaoke' → ຄົ້ນຫາ",
             "ຜົນລັບສະແດງສະເພາະ karaoke version", "", ""],
        ]
    )

    doc.add_page_break()

    # ===== TC-04 =====
    doc.add_heading("6. TC-04: ຈັດການຄິວເພງ", level=1)
    add_styled_table(doc,
        ["ID", "ລາຍການທົດສອບ", "ຂັ້ນຕອນ", "ຜົນທີ່ຄາດຫວັງ", "ຜົນຈິງ", "ຜ່ານ?"],
        [
            ["4.1", "ເພີ່ມເພງເຂົ້າຄິວ",
             "Browse → ກົດ + ທີ່ເພງ",
             "ເພງເພີ່ມໃນຄິວ, ສະແດງໃນ Queue", "", ""],
            ["4.2", "ຍ້າຍເພງຂຶ້ນ",
             "ໜ້າ Queue → ກົດ ຍ້າຍຂຶ້ນ",
             "ເພງຍ້າຍຂຶ້ນ 1 ຕຳແໜ່ງ", "", ""],
            ["4.3", "ຍ້าຍເພງລົງ",
             "ໜ້າ Queue → ກົດ ຍ້າຍລົງ",
             "ເພງຍ້າຍລົງ 1 ຕຳແໜ່ງ", "", ""],
            ["4.4", "ລຶບເພງຈາກຄິວ",
             "ກົດ ລຶບ ທີ່ເພງ",
             "ເພງຖືກເອົາອອກຈາກຄິວ", "", ""],
            ["4.5", "Drag & Drop",
             "ລາກເພງໄປຕຳແໜ່ງໃໝ່",
             "ລຳດັບຄິວປ່ຽນຕາມ", "", ""],
            ["4.6", "ເພີ່ມ 3 ເພງສຸ່ມ",
             "ກົດ 'Add 3 Random Songs'",
             "3 ເພງສຸ່ມເພີ່ມເຂົ້າຄິວ", "", ""],
            ["4.7", "ລ້າງຄິວທັງໝົດ",
             "ກົດ Clear All → ຢືນຢັນ",
             "ຄິວຫວ່າງ", "", ""],
            ["4.8", "ຄິວຫວ່າງ",
             "ລ້າງຄິວ → ເບິ່ງໜ້າ Queue",
             "ສະແດງ 'Queue is empty'", "", ""],
        ]
    )

    doc.add_page_break()

    # ===== TC-05 =====
    doc.add_heading("7. TC-05: ເບິ່ງ ແລະ ຈັດການເພງ (Browse)", level=1)
    add_styled_table(doc,
        ["ID", "ລາຍການທົດສອບ", "ຂັ້ນຕອນ", "ຜົນທີ່ຄາດຫວັງ", "ຜົນຈິງ", "ຜ່ານ?"],
        [
            ["5.1", "ສະແດງລາຍການເພງ",
             "ໄປໜ້າ Browse",
             "ສະແດງເພງທັງໝົດພ້ອມເລກລຳດັບ", "", ""],
            ["5.2", "ຈັດລຽງ A-Z",
             "ກົດ Alphabetical sort",
             "ເພງຈັດລຽງຕາມຕົວອັກສອນ", "", ""],
            ["5.3", "ຈັດລຽງຕາມວັນທີ",
             "ກົດ Date sort",
             "ເພງຈັດລຽງຕາມວັນທີ download", "", ""],
            ["5.4", "ກັ່ນຕອງດ້ວຍຕົວອັກສອນ",
             "ກົດ 'S' ໃນແຖບ A-Z",
             "ສະແດງສະເພາະເພງທີ່ຂຶ້ນຕົ້ນດ້ວຍ S", "", ""],
            ["5.5", "ເພີ່ມເພງເຂົ້າຄິວ",
             "ກົດ + ທີ່ເພງ",
             "ເພງເພີ່ມໃນຄິວ", "", ""],
            ["5.6", "ແກ້ໄຊຊື່ເພງ (Admin)",
             "ກົດ pencil icon → ປ່ຽນຊື່ → Save",
             "ຊື່ເພງປ່ຽນ", "", ""],
        ]
    )

    # ===== TC-06 =====
    doc.add_heading("8. TC-06: ການໃຊ້ງານຜ່ານມືຖື", level=1)
    add_styled_table(doc,
        ["ID", "ລາຍການທົດສອບ", "ຂັ້ນຕອນ", "ຜົນທີ່ຄາດຫວັງ", "ຜົນຈິງ", "ຜ່ານ?"],
        [
            ["6.1", "ເຂົ້າເຖິງຈາກມືຖື",
             "ສະແກນ QR Code ຫຼື ພິມ IP:port",
             "ເປີດໜ້າ app ໄດ້ປົກກະຕິ", "", ""],
            ["6.2", "Responsive layout",
             "ເປີດທຸກໜ້າໃນມືຖື",
             "UI ປັບຕາມໜ້າຈໍ, ມີ hamburger menu", "", ""],
            ["6.3", "ຄົ້ນຫາ ແລະ ເພີ່ມຄິວ",
             "ຄົ້ນຫາເພງ → ເພີ່ມຄິວ ຈາກມືຖື",
             "ເພງເພີ່ມໃນຄິວ, TV ອັບເດດ real-time", "", ""],
            ["6.4", "ໃສ່ຊື່ຜູ້ຮ້ອງ",
             "ກົດ user icon → ໃສ່ຊື່",
             "ຊື່ບັນທຶກ ແລະ ສະແດງໃນຄິວ", "", ""],
            ["6.5", "ຫຼາຍຄົນໃຊ້ພ້ອມກັນ",
             "2-3 ມືຖືເຂົ້າພ້ອ�ກັນ",
             "ທຸກຄົນເຫັນຄິວອັບເດດ real-time", "", ""],
        ]
    )

    doc.add_page_break()

    # ===== TC-07 =====
    doc.add_heading("9. TC-07: ລະບົບ Admin", level=1)
    add_styled_table(doc,
        ["ID", "ລາຍການທົດສອບ", "ຂັ້ນຕອນ", "ຜົນທີ່ຄາດຫວັງ", "ຜົນຈິງ", "ຜ່ານ?"],
        [
            ["7.1", "Login admin",
             "ໄປໜ້າ Info → ກົດ Login → ໃສ່ password",
             "Login ສຳເລັດ, ເຫັນປຸ່ມ admin", "", ""],
            ["7.2", "Login ລະຫັດຜິດ",
             "ໃສ່ password ຜິດ",
             "ແຈ້ງ error, ບໍ່ໃຫ້ເຂົ້າ", "", ""],
            ["7.3", "Logout",
             "ກົດ Logout",
             "ປຸ່ມ admin ຫາຍໄປ", "", ""],
            ["7.4", "ບໍ່ມີ password (default)",
             "ແລ່ນ app ບໍ່ໃສ່ --admin-password",
             "ທຸກຄົນໃຊ້ຟີເຈີ admin ໄດ້", "", ""],
            ["7.5", "ການປ້ອງກັນ",
             "ບໍ່ login → ລອງລຶບເພງຈາກຄິວ",
             "ບໍ່ອະນຸຍາດ, ປຸ່ມ disabled", "", ""],
        ]
    )

    # ===== TC-08 =====
    doc.add_heading("10. TC-08: ໜ້າ Info ແລະ ການຕັ້ງຄ່າ", level=1)
    add_styled_table(doc,
        ["ID", "ລາຍການທົດສອບ", "ຂັ້ນຕອນ", "ຜົນທີ່ຄາດຫວັງ", "ຜົນຈິງ", "ຜ່ານ?"],
        [
            ["8.1", "ສະແດງ QR Code",
             "ເປີດໜ້າ Info",
             "ສະແດງ QR Code ແລະ URL", "", ""],
            ["8.2", "ສະແດງຂໍ້ມູນລະບົບ",
             "ເບິ່ງ System Info",
             "ສະແດງ CPU, RAM, Disk, yt-dlp version", "", ""],
            ["8.3", "ປ່ຽນພາສາ",
             "ເລືອກພາສາໃໝ່ຈາກ dropdown",
             "Interface ປ່ຽນເປັນພາສາທີ່ເລືອກ", "", ""],
            ["8.4", "Refresh Database",
             "ກົດ Refresh",
             "ສະແກນ folder ເພງໃໝ່", "", ""],
            ["8.5", "Update yt-dlp",
             "ກົດ Update",
             "yt-dlp ອັບເດດ (ຫຼື ແຈ້ງວ່າໃໝ່ແລ້ວ)", "", ""],
            ["8.6", "Volume Normalization",
             "ເປີດ toggle → ຫຼິ້ນຫຼາຍເພງ",
             "Volume ເທົ່າກັນທຸກເພງ", "", ""],
        ]
    )

    doc.add_page_break()

    # ===== TC-09 =====
    doc.add_heading("11. TC-09: ການທົດສອບ Edge Cases", level=1)
    add_styled_table(doc,
        ["ID", "ລາຍການທົດສອບ", "ຂັ້ນຕອນ", "ຜົນທີ່ຄາດຫວັງ", "ຜົນຈິງ", "ຜ່ານ?"],
        [
            ["9.1", "ຄິວຫວ່າງ + Skip",
             "ຄິວຫວ່າງ → ກົດ Skip",
             "ບໍ່ crash, ສະແດງ splash screen", "", ""],
            ["9.2", "ດາວໂຫຼດເພງຊ້ຳ",
             "ດາວໂຫຼດເພງທີ່ມີຢູ່ແລ້ວ",
             "ແຈ້ງວ່າມີແລ້ວ ຫຼື overwrite", "", ""],
            ["9.3", "ຄົ້ນຫາຫວ່າງ",
             "ກົດ Search ໂດຍບໍ່ໃສ່ຂໍ້ຄວາມ",
             "ບໍ່ crash, ສະແດງ warning", "", ""],
            ["9.4", "URL ບໍ່ຖືກຕ້ອງ",
             "ວາງ URL ທີ່ບໍ່ແມ່ນ YouTube → Download",
             "ແຈ້ງ error ຢ່າງເໝາະສົມ", "", ""],
            ["9.5", "ຕັດ internet ກະທັນຫັນ",
             "ດາວໂຫຼດເພງ → ປິດ WiFi",
             "ແຈ້ງ error, app ບໍ່ crash", "", ""],
            ["9.6", "ເປີດຫຼາຍ browser tabs",
             "ເປີດ 3+ tabs ພ້ອມກັນ",
             "ທຸກ tab ອັບເດດ real-time", "", ""],
            ["9.7", "ຊື່ເພງມີອັກສອນພິເສດ",
             "ດາວໂຫຼດເພງຊື່ມີ /, \\, ?, #",
             "ດາວໂຫຼດສຳເລັດ, ຊື່ຖືກ sanitize", "", ""],
            ["9.8", "Folder ເພງຫວ່າງ",
             "ແລ່ນ app ກັບ folder ຫວ່າງ",
             "App ເປີດປົກກະຕິ, Browse ສະແດງ empty", "", ""],
        ]
    )

    doc.add_page_break()

    # ===== 12. ສະຫຼຸບ =====
    doc.add_heading("12. ສະຫຼຸບ ແລະ ເກນການຜ່ານ", level=1)

    doc.add_heading("ສະຫຼຸບ Test Cases", level=2)
    add_styled_table(doc,
        ["ໝວດ", "ຈຳນວນ test cases"],
        [
            ["TC-01: ຕິດຕັ້ງ ແລະ ເປີດ App", "5"],
            ["TC-02: ຄວບຄຸມການຫຼິ້ນ", "10"],
            ["TC-03: ຄົ້ນຫາ ແລະ ດາວໂຫຼດ", "7"],
            ["TC-04: ຈັດການຄິວ", "8"],
            ["TC-05: Browse ເພງ", "6"],
            ["TC-06: ມືຖື", "5"],
            ["TC-07: Admin", "5"],
            ["TC-08: Info ແລະ ຕັ້ງຄ່າ", "6"],
            ["TC-09: Edge Cases", "8"],
            ["ລວມທັງໝົດ", "60"],
        ],
        header_color="1a5276"
    )

    doc.add_heading("ເກນການຜ່ານ", level=2)
    criteria = [
        "Critical tests (TC-01, TC-02, TC-03): ຕ້ອງຜ່ານ 100%",
        "Major tests (TC-04 ~ TC-08): ຕ້ອງຜ່ານ 90% ຂຶ້ນໄປ",
        "Edge cases (TC-09): ຕ້ອງຜ່ານ 75% ຂຶ້ນໄປ",
        "ລວມທັງໝົດ: ຕ້ອງຜ່ານ 90% ຂຶ້ນໄປ ຈຶ່ງຖືວ່າພ້ອມນຳໃຊ້",
    ]
    for c in criteria:
        doc.add_paragraph(c, style='List Bullet')

    doc.add_heading("ການລາຍງານ Bug", level=2)
    doc.add_paragraph("ຖ້າພົບ bug ໃຫ້ບັນທຶກ:")
    bug_info = [
        "ລະຫັດ test case (ເຊັ່ນ TC-02, ID 2.5)",
        "ຂັ້ນຕອນທີ່ເຮັດ",
        "ຜົນທີ່ຄາດຫວັງ vs ຜົນຈິງ",
        "Screenshot (ຖ້າມີ)",
        "Browser ແລະ ເວີຊັນ",
    ]
    for item in bug_info:
        doc.add_paragraph(item, style='List Bullet')

    # Save
    output_path = os.path.join(os.path.dirname(__file__), "TestGuide_OpenHomeKaraoke.docx")
    doc.save(output_path)
    print(f"Test Guide saved: {output_path}")
    return output_path


if __name__ == "__main__":
    print("Creating guides...")
    create_user_guide()
    create_test_guide()
    print("Done!")
